# module fileuploader

# system
import os, io

# loaders
import docx
from PyPDF2 import PdfReader
from odf import text as odfText
from odf.opendocument import load as odtload
from streamlit.runtime.uploaded_file_manager import UploadedFile
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


def fu_make_upload(file: UploadedFile | str) -> tuple[io.BytesIO, str, str]:
    """Converts file to tuple pair of file data and type, and loads the file to memory

    Args:
        file (UploadedFile | str): uploaded file from uploader or filepath

    Returns:
        tuple[io.BytesIO, str, str]: file data, file type, file name
    """
    if isinstance(file, UploadedFile):
        return (io.BytesIO(file.read()), file.type, file.name)
    else:
        with open(file, 'rb') as f:
            data = f.read()
            type = file.split('.')[-1]
        return (io.BytesIO(data), fu_get_encoded_type(type), os.path.basename(file))


def fu_get_encoded_type(ext: str) -> str:
    """Returns official document type name based off its extension

    Args:
        ext (str): file extension

    Raises:
        Exception: unsupported document type

    Returns:
        str: document type
    
    Note:
        supported document types: ['pdf', 'odt', 'docx', 'txt']
    """
    if ext == 'pdf': return 'application/pdf'
    elif ext == 'odt': return 'application/vnd.oasis.opendocument.text'
    elif ext == 'docx': return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif ext == 'txt': return 'text/plain'
    else: raise Exception('Unsupported document type!')


def fu_get_content(file: io.BytesIO, type: str, source: str = '') -> list[Document]:
    """Reads file contents based in its text

    Args:
        file (io.BytesIO): uploaded file object
        type (str): document type
        source (str, optional): filename. Defaults to ''.

    Returns:
        list[Document]: list of documents(page_content, metadata={'source', 'page'})
    """
    if type == 'application/pdf': return fu_get_content_pdf(file, source)
    elif type == 'application/vnd.oasis.opendocument.text': return fu_get_content_odt(file, source)
    elif type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document': return fu_get_content_docx(file, source)
    elif type == 'text/plain': return fu_get_content_txt(file, source)
    else: raise Exception('Unsupported document type!')


def fu_get_content_txt(file: io.BytesIO, source: str = '') -> list[Document]:
    """Reads TXT file contents

    Args:
        file (io.BytesIO): uploaded file object
        source (str, optional): filename. Defaults to ''.

    Returns:
        list[Document]: list of documents(page_content, metadata={'source', 'page'})
    """
    documents = [Document(
        page_content=file.read().decode(errors='replace'),
        metadata={
            'source': source,
            'page': 0,
        }
    )]
    return documents


def fu_get_content_pdf(file: io.BytesIO, source: str = '') -> list[Document]:
    """Reads PDF file contents

    Args:
        file (io.BytesIO): uploaded file object
        source (str, optional): filename. Defaults to ''.

    Returns:
        list[Document]: list of documents(page_content, metadata={'source', 'page'})
    """
    documents = []
    loader = PdfReader(file)
    for i, chunk in enumerate(loader.pages):
        documents.append(Document(
            page_content=chunk.extract_text(),
            metadata={
                'source': source,
                'page': i,
            }
        ))
    return documents


def fu_get_content_odt(file: io.BytesIO, source: str = '') -> list[Document]:
    """Reads ODT file contents

    Args:
        file (io.BytesIO): uploaded file object
        source (str, optional): filename. Defaults to ''.

    Returns:
        list[Document]: list of documents(page_content, metadata={'source', 'page'})
    """    
    documents = []
    loader = odtload(file)
    for i, chunk in enumerate(loader.getElementsByType(odfText.P)):
        documents.append(Document(
            page_content=chunk.firstChild.data,
            metadata={
                'source': source,
                'page': i,
            }
        ))
    return documents


def fu_get_content_docx(file: io.BytesIO, source: str = '') -> list[Document]:
    """Reads DOCX file contents

    Args:
        file (io.BytesIO): uploaded file object
        source (str, optional): filename. Defaults to ''.

    Returns:
        list[Document]: list of documents(page_content, metadata={'source', 'page'})
    """
    documents = []
    loader = docx.Document(file)
    for i, chunk in enumerate(loader.paragraphs): 
        documents.append(Document(
            page_content=chunk.text,
            metadata={
                'source': source,
                'page': i,
            }
        ))
    return documents


def fu_split_documents(documents: list[Document], chunk_size: int = 800, chunk_overlap: int = 80) -> list[Document]:
    """Split documents into smaller chunks of specified size

    Args:
        documents (list[Document]): list of documents
        chunk_size (int, optional): chunk size split. Defaults to 800.
        chunk_overlap (int, optional): chunk size overlap. Defaults to 80.

    Returns:
        list[Document]: list of documents of the approximate specified size
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=False,
    )
    return text_splitter.split_documents(documents)


def fu_calculate_document_ids(documents: list[Document]) -> list[Document]:
    """Creates unique document IDs like 'source:page_num:document_id' ==> 'file.pdf:6:2'

    Args:
        documents (list[Document]): documents

    Returns:
        list[Document]: documents with unique id
    """
    last_page_id = None
    current_chunk_index = 0
    for document in documents:
        source = document.metadata.get("source")
        page = document.metadata.get("page")
        current_page_id = f"{source}:{page}"

        # if the page ID is the same as the last one, increment the index.
        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        # calculate the document ID.
        document_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id

        # add it to the page meta-data.
        document.metadata["id"] = document_id

    return documents
